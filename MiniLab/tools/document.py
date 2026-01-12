"""
Document Generation Tool.

Provides document generation capabilities:
- doc.docx: Generate DOCX documents
- doc.pdf: Generate PDF documents
- doc.to_docx: Convert markdown to DOCX
- doc.to_pdf: Convert markdown to PDF

Requires optional dependencies:
- python-docx for DOCX generation
- pypandoc or weasyprint for PDF generation
"""

from __future__ import annotations

import os
import re
import subprocess
from pathlib import Path
from typing import Any, Optional, Union
from pydantic import Field

from .base import Tool, ToolInput, ToolOutput, ToolError
from ..security import PathGuard, AccessDenied
from ..utils import console


class CreateDocxInput(ToolInput):
    """Input for creating a DOCX document."""
    output_path: str = Field(..., description="Path for the output DOCX file")
    title: str = Field("", description="Document title")
    content: str = Field(..., description="Document content (markdown or plain text)")
    author: str = Field("MiniLab", description="Document author")
    template_path: Optional[str] = Field(None, description="Optional template DOCX path")


class CreatePdfInput(ToolInput):
    """Input for creating a PDF document."""
    output_path: str = Field(..., description="Path for the output PDF file")
    title: str = Field("", description="Document title")
    content: str = Field(..., description="Document content (markdown or plain text)")
    author: str = Field("MiniLab", description="Document author")


class MarkdownToDocxInput(ToolInput):
    """Input for converting markdown to DOCX."""
    input_path: str = Field(..., description="Path to the markdown file")
    output_path: str = Field(..., description="Path for the output DOCX file")
    template_path: Optional[str] = Field(None, description="Optional template DOCX path")


class MarkdownToPdfInput(ToolInput):
    """Input for converting markdown to PDF."""
    input_path: str = Field(..., description="Path to the markdown file")
    output_path: str = Field(..., description="Path for the output PDF file")


class DocumentOutput(ToolOutput):
    """Output for document operations."""
    path: Optional[str] = None
    size_bytes: Optional[int] = None
    page_count: Optional[int] = None


class DocumentTool(Tool):
    """
    Tool for generating DOCX and PDF documents.
    
    Supports:
    - Creating documents from content
    - Converting markdown to DOCX/PDF
    - Using templates
    
    Requires system dependencies:
    - pandoc (for markdown conversion)
    - Optional: python-docx, pypandoc
    """
    
    name = "document"
    description = "Generate DOCX and PDF documents"
    
    def get_actions(self) -> dict[str, str]:
        return {
            "create_docx": "Create a DOCX document from content",
            "create_pdf": "Create a PDF document from content",
            "markdown_to_docx": "Convert markdown file to DOCX",
            "markdown_to_pdf": "Convert markdown file to PDF",
        }
    
    def get_input_schema(self, action: str) -> type[ToolInput]:
        schemas = {
            "create_docx": CreateDocxInput,
            "create_pdf": CreatePdfInput,
            "markdown_to_docx": MarkdownToDocxInput,
            "markdown_to_pdf": MarkdownToPdfInput,
        }
        return schemas.get(action, ToolInput)
    
    async def execute(self, action: str, params: dict[str, Any]) -> ToolOutput:
        """Execute a document action."""
        schema = self.get_input_schema(action)
        validated = schema(**params)
        
        try:
            if action == "create_docx":
                return await self._create_docx(validated)
            elif action == "create_pdf":
                return await self._create_pdf(validated)
            elif action == "markdown_to_docx":
                return await self._md_to_docx(validated)
            elif action == "markdown_to_pdf":
                return await self._md_to_pdf(validated)
            else:
                return DocumentOutput(
                    success=False,
                    error=f"Unknown action: {action}",
                )
        except Exception as e:
            return DocumentOutput(success=False, error=str(e))
    
    async def _create_docx(self, params: CreateDocxInput) -> DocumentOutput:
        """Create a DOCX document from content."""
        from ..security import PathGuard
        
        output_path = Path(params.output_path)
        
        # Check write permission
        guard = PathGuard.get_instance()
        if not guard.can_write(output_path):
            raise AccessDenied(str(output_path), "write")
        
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            # Try using python-docx
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            if params.template_path and Path(params.template_path).exists():
                doc = Document(params.template_path)
            else:
                doc = Document()
            
            # Add title
            if params.title:
                title = doc.add_heading(params.title, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content - basic markdown parsing
            content = params.content
            current_para = []
            
            for line in content.split('\n'):
                line = line.strip()
                
                # Headings
                if line.startswith('# '):
                    if current_para:
                        doc.add_paragraph(' '.join(current_para))
                        current_para = []
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    if current_para:
                        doc.add_paragraph(' '.join(current_para))
                        current_para = []
                    doc.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    if current_para:
                        doc.add_paragraph(' '.join(current_para))
                        current_para = []
                    doc.add_heading(line[4:], 3)
                # Bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    if current_para:
                        doc.add_paragraph(' '.join(current_para))
                        current_para = []
                    doc.add_paragraph(line[2:], style='List Bullet')
                # Numbered lists
                elif re.match(r'^\d+\.\s', line):
                    if current_para:
                        doc.add_paragraph(' '.join(current_para))
                        current_para = []
                    doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
                # Empty line = paragraph break
                elif not line:
                    if current_para:
                        doc.add_paragraph(' '.join(current_para))
                        current_para = []
                else:
                    current_para.append(line)
            
            # Add remaining content
            if current_para:
                doc.add_paragraph(' '.join(current_para))
            
            # Save
            doc.save(str(output_path))
            
            size = output_path.stat().st_size
            
            return DocumentOutput(
                success=True,
                path=str(output_path),
                size_bytes=size,
            )
            
        except ImportError:
            # Fall back to pandoc
            return await self._create_with_pandoc(
                params.content, output_path, "docx", params.title
            )
    
    async def _create_pdf(self, params: CreatePdfInput) -> DocumentOutput:
        """Create a PDF document from content."""
        from ..security import PathGuard
        
        output_path = Path(params.output_path)
        
        # Check write permission
        guard = PathGuard.get_instance()
        if not guard.can_write(output_path):
            raise AccessDenied(str(output_path), "write")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Use pandoc for PDF generation
        return await self._create_with_pandoc(
            params.content, output_path, "pdf", params.title
        )
    
    async def _md_to_docx(self, params: MarkdownToDocxInput) -> DocumentOutput:
        """Convert markdown file to DOCX."""
        from ..security import PathGuard
        
        input_path = Path(params.input_path)
        output_path = Path(params.output_path)
        
        guard = PathGuard.get_instance()
        if not guard.can_read(input_path):
            raise AccessDenied(str(input_path), "read")
        if not guard.can_write(output_path):
            raise AccessDenied(str(output_path), "write")
        
        if not input_path.exists():
            return DocumentOutput(
                success=False,
                error=f"Input file not found: {input_path}",
            )
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Use pandoc
        cmd = ["pandoc", str(input_path), "-o", str(output_path)]
        
        if params.template_path and Path(params.template_path).exists():
            cmd.extend(["--reference-doc", params.template_path])
        
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
            )
            
            if result.returncode != 0:
                return DocumentOutput(
                    success=False,
                    error=f"pandoc error: {result.stderr}",
                )
            
            size = output_path.stat().st_size
            
            return DocumentOutput(
                success=True,
                path=str(output_path),
                size_bytes=size,
            )
            
        except FileNotFoundError:
            return DocumentOutput(
                success=False,
                error="pandoc not found. Install with: brew install pandoc",
            )
        except subprocess.TimeoutExpired:
            return DocumentOutput(
                success=False,
                error="pandoc conversion timed out",
            )
    
    async def _md_to_pdf(self, params: MarkdownToPdfInput) -> DocumentOutput:
        """Convert markdown file to PDF."""
        from ..security import PathGuard
        
        input_path = Path(params.input_path)
        output_path = Path(params.output_path)
        
        guard = PathGuard.get_instance()
        if not guard.can_read(input_path):
            raise AccessDenied(str(input_path), "read")
        if not guard.can_write(output_path):
            raise AccessDenied(str(output_path), "write")
        
        if not input_path.exists():
            return DocumentOutput(
                success=False,
                error=f"Input file not found: {input_path}",
            )
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Use pandoc with PDF engine
        cmd = [
            "pandoc", str(input_path),
            "-o", str(output_path),
            "--pdf-engine=xelatex",  # or pdflatex, lualatex
        ]
        
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
            )
            
            if result.returncode != 0:
                # Try without pdf-engine specification
                cmd = ["pandoc", str(input_path), "-o", str(output_path)]
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                
                if result.returncode != 0:
                    return DocumentOutput(
                        success=False,
                        error=f"pandoc error: {result.stderr}",
                    )
            
            size = output_path.stat().st_size
            
            return DocumentOutput(
                success=True,
                path=str(output_path),
                size_bytes=size,
            )
            
        except FileNotFoundError:
            return DocumentOutput(
                success=False,
                error="pandoc not found. Install with: brew install pandoc",
            )
        except subprocess.TimeoutExpired:
            return DocumentOutput(
                success=False,
                error="PDF conversion timed out",
            )
    
    async def _create_with_pandoc(
        self,
        content: str,
        output_path: Path,
        format: str,
        title: Optional[str] = None,
    ) -> DocumentOutput:
        """Create document using pandoc."""
        import tempfile
        
        # Write content to temp file
        with tempfile.NamedTemporaryFile(
            mode='w',
            suffix='.md',
            delete=False
        ) as f:
            if title:
                f.write(f"# {title}\n\n")
            f.write(content)
            temp_path = f.name
        
        try:
            cmd = ["pandoc", temp_path, "-o", str(output_path)]
            
            if format == "pdf":
                cmd.append("--pdf-engine=xelatex")
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
            )
            
            if result.returncode != 0:
                # Retry without pdf-engine
                if format == "pdf":
                    cmd = ["pandoc", temp_path, "-o", str(output_path)]
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        timeout=120,
                    )
                
                if result.returncode != 0:
                    return DocumentOutput(
                        success=False,
                        error=f"pandoc error: {result.stderr}",
                    )
            
            size = output_path.stat().st_size
            
            return DocumentOutput(
                success=True,
                path=str(output_path),
                size_bytes=size,
            )
            
        except FileNotFoundError:
            return DocumentOutput(
                success=False,
                error="pandoc not found. Install with: brew install pandoc",
            )
        finally:
            # Clean up temp file
            Path(temp_path).unlink(missing_ok=True)
